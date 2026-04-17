"""
resume_parser.py
----------------
Handles extracting text from PDF and DOCX resume files,
extracting candidate name, phone, email, and
cleaning the raw text for further processing.
"""

import re
import os
from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(filepath):
    """
    Extract all text from a PDF file.
    
    Args:
        filepath (str): Path to the PDF file.
    
    Returns:
        str: Extracted text from all pages.
    """
    text = ""
    try:
        reader = PdfReader(filepath)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"[ERROR] Failed to read PDF: {e}")
    return text


def extract_text_from_docx(filepath):
    """
    Extract all text from a DOCX file, including tables.
    
    Args:
        filepath (str): Path to the DOCX file.
    
    Returns:
        str: Extracted text from all paragraphs and tables.
    """
    text = ""
    try:
        doc = Document(filepath)

        # Extract paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"

        # Extract text from tables (e.g. skills in tabular format)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
    except Exception as e:
        print(f"[ERROR] Failed to read DOCX: {e}")
    return text


def extract_text(filepath):
    """
    Detect file type and extract text accordingly.
    
    Args:
        filepath (str): Path to a PDF or DOCX file.
    
    Returns:
        str: Extracted raw text.
    
    Raises:
        ValueError: If the file format is not supported.
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Use PDF or DOCX.")


def extract_email(raw_text):
    """
    Extract email addresses from resume text.
    
    Args:
        raw_text (str): The raw extracted text.
    
    Returns:
        str: First email found, or empty string.
    """
    pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    emails = re.findall(pattern, raw_text)
    return emails[0] if emails else ""


def extract_phone(raw_text):
    """
    Extract phone numbers from resume text.
    Supports Indian (+91), US (+1), and international formats.
    
    Args:
        raw_text (str): The raw extracted text.
    
    Returns:
        str: First phone number found, or empty string.
    """
    patterns = [
        r'(?:\+91[\s-]?)?[6-9]\d{4}[\s-]?\d{5}',     # Indian: +91 98765 43210
        r'(?:\+1[\s-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}',  # US: (555) 123-4567
        r'\+?\d{1,3}[\s.-]?\d{3,4}[\s.-]?\d{3,4}[\s.-]?\d{0,4}',  # International
    ]
    for pattern in patterns:
        phones = re.findall(pattern, raw_text)
        if phones:
            # Return the longest match (most likely a real phone number)
            return max(phones, key=len).strip()
    return ""


def extract_candidate_name(raw_text):
    """
    Attempt to extract the candidate's name from the first few lines of the resume.
    Usually the first non-empty, non-email, non-phone line is the name.
    
    Args:
        raw_text (str): The raw extracted text.
    
    Returns:
        str: Candidate name or empty string.
    """
    lines = raw_text.strip().split("\n")
    for line in lines[:5]:  # Check first 5 lines
        line = line.strip()
        if not line:
            continue
        # Skip if it's an email
        if re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', line):
            continue
        # Skip if it's a phone number
        if re.search(r'[\+]?\d[\d\s\-\(\)]{7,}', line):
            continue
        # Skip common headers
        skip_words = ["resume", "curriculum vitae", "cv", "objective", "summary", "profile"]
        if line.lower().strip() in skip_words:
            continue
        # If line is short and contains mostly letters, it's likely a name
        if len(line) < 60 and re.match(r'^[A-Za-z\s\.\-]+$', line):
            return line.title()
    return ""


def clean_text(raw_text):
    """
    Clean and normalize extracted text.
    
    Steps:
        1. Convert to lowercase
        2. Remove URLs
        3. Remove email addresses
        4. Remove special characters (keep letters, numbers, spaces, and +, #, -, ., /)
        5. Collapse multiple spaces into one
        6. Strip leading/trailing whitespace
    
    Args:
        raw_text (str): The raw extracted text.
    
    Returns:
        str: Cleaned text ready for NLP processing.
    """
    text = raw_text.lower()

    # Remove URLs
    text = re.sub(r"http\S+|www\.\S+", "", text)

    # Remove email addresses
    text = re.sub(r"\S+@\S+\.\S+", "", text)

    # Remove special characters but keep letters, numbers, spaces, and specific symbols (+, #, -, ., /)
    text = re.sub(r"[^a-z0-9\s\+\#\-\.\/]", " ", text)

    # Collapse multiple spaces
    text = re.sub(r"\s+", " ", text)

    return text.strip()